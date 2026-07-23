#!/usr/bin/env python3
# Build an editable Word (.docx) of the Burton Inn lease for redlining.
# Combines Part 1 (burton-inn/terms-sheet.md — Terms Sheet + Exhibits A-D) and
# the v1.5 Standard Lease Terms + Definitions (lease/lease.md, Parts 2 & 3,
# reused verbatim) into one document, mirroring the merged lease.pdf. Uses real
# Word heading styles so the document is navigable and Track Changes works
# cleanly. Plain python-docx; absolute paths.
#
# Note: Parts 2 & 3 are the vetted v1.5 master form — the intent is to redline
# Part 1 (Terms Sheet + Exhibit D amendment) only and leave the form unchanged.

import os, re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BDIR = ROOT + '/burton-inn'
INK = RGBColor(0x1E, 0x31, 0x28)   # match the PDF's deep evergreen headings


def add_runs(paragraph, text):
    """Render inline **bold** and *italic* markdown into Word runs."""
    for i, chunk in enumerate(re.split(r'\*\*(.+?)\*\*', text)):
        if i % 2 == 1:
            paragraph.add_run(chunk).bold = True
        else:
            for j, sub in enumerate(re.split(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', chunk)):
                if not sub:
                    continue
                run = paragraph.add_run(sub)
                if j % 2 == 1:
                    run.italic = True


def emit(doc, md, first):
    """Append one markdown document's blocks to the Word doc."""
    lines = md.split('\n')
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if s == '':
            i += 1; continue
        if s == '---':                      # section divider — skip (headings carry structure)
            i += 1; continue
        if s.startswith('# '):              # document title
            if not first:
                doc.add_page_break()
            first = False
            doc.add_heading(s[2:].strip(), level=0)
            i += 1; continue
        if s.startswith('## '):
            doc.add_heading(s[3:].strip(), level=1)
            i += 1; continue
        if s.startswith('### '):
            doc.add_heading(s[4:].strip(), level=2)
            i += 1; continue
        if s.startswith('> '):              # applicability / note blockquote
            p = doc.add_paragraph(style='Intense Quote')
            add_runs(p, s[2:].strip())
            i += 1; continue
        if s.startswith('- '):              # bullet list
            while i < len(lines) and lines[i].strip().startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                add_runs(p, lines[i].strip()[2:])
                i += 1
            continue
        p = doc.add_paragraph()             # ordinary paragraph (incl. bold-lead fields)
        add_runs(p, s)
        i += 1
    return first


def style_document(doc):
    """Legal-document look: Times New Roman body, black serif headings."""
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(11)
    for name in ('Title', 'Heading 1', 'Heading 2'):
        st = doc.styles[name]
        st.font.name = 'Times New Roman'
        st.font.color.rgb = INK
    # Centre the title block's header lines are left as-is; headings navigable via the styles.


if __name__ == '__main__':
    with open(BDIR + '/terms-sheet.md') as f:
        ts = f.read()
    with open(ROOT + '/lease/lease.md') as f:
        form = f.read()          # Parts 2 & 3, verbatim v1.5

    doc = Document()
    style_document(doc)
    first = emit(doc, ts, first=True)
    emit(doc, form, first=first)

    out = BDIR + '/lease.docx'
    doc.save(out)
    print('wrote', out)
